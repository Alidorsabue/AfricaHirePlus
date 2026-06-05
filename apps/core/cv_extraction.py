"""
cv_extraction v2 — Extraction robuste multi-format depuis les CV.

Formats supportés : PDF (natif + OCR fallback), Word (.docx/.doc),
                    texte brut (.txt), RTF (.rtf), ODT (.odt),
                    images CV (.jpg, .jpeg, .png, .webp, .tiff, .bmp).

Pipeline :
  1. Détection du format (content_type + extension + magic bytes)
  2. Extraction primaire par moteur adapté
  3. Fallback automatique si texte insuffisant (PDF natif → pdfminer → OCR)
  4. Nettoyage et normalisation du texte
  5. Détection qualité + métadonnées (nb pages, score qualité, méthode)

Utilisé en amont de ml/feature_engineering.py pour remplir candidate.raw_cv_text.
Compatible Django UploadedFile, fichiers-like, bytes bruts.
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from enum import Enum
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────
# CONSTANTES
# ─────────────────────────────────────────────────────────────

# Seuil de texte « suffisant » (nb caractères) pour considérer l'extraction réussie
MINIMUM_TEXT_LENGTH = 100

# Seuil en dessous duquel on tente l'OCR sur un PDF
PDF_OCR_FALLBACK_THRESHOLD = 50

# Magic bytes pour détection de format fiable (indépendant du content_type)
MAGIC_BYTES: list[tuple[bytes, str]] = [
    (b"%PDF",                           "pdf"),
    (b"PK\x03\x04",                     "docx"),   # ZIP → docx/odt
    (b"\xd0\xcf\x11\xe0",              "doc"),    # OLE2 → doc/xls
    (b"{\\rtf",                         "rtf"),
    (b"\xff\xd8\xff",                   "jpg"),
    (b"\x89PNG\r\n\x1a\n",            "png"),
    (b"RIFF",                           "webp"),
    (b"II*\x00",                        "tiff"),
    (b"MM\x00*",                        "tiff"),
    (b"BM",                             "bmp"),
]

DOCX_MIME = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.oasis.opendocument.text",
}
IMAGE_MIME = {"image/jpeg", "image/png", "image/webp", "image/tiff", "image/bmp"}


class ExtractionMethod(str, Enum):
    PYPDF       = "pypdf"
    PDFMINER    = "pdfminer"
    OCR         = "ocr"
    DOCX        = "docx"
    DOC         = "doc"
    ODT         = "odt"
    RTF         = "rtf"
    PLAIN_TEXT  = "plain_text"
    IMAGE_OCR   = "image_ocr"
    NONE        = "none"


@dataclass
class ExtractionResult:
    text: str = ""
    method: ExtractionMethod = ExtractionMethod.NONE
    page_count: int = 0
    quality_score: float = 0.0   # 0.0–1.0 (densité de texte utile)
    ocr_used: bool = False
    warnings: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def is_sufficient(self) -> bool:
        return len(self.text.strip()) >= MINIMUM_TEXT_LENGTH

    def to_dict(self) -> dict:
        return {
            "method": self.method.value,
            "page_count": self.page_count,
            "quality_score": round(self.quality_score, 3),
            "ocr_used": self.ocr_used,
            "char_count": len(self.text),
            "warnings": self.warnings,
        }


# ─────────────────────────────────────────────────────────────
# DÉTECTION DE FORMAT
# ─────────────────────────────────────────────────────────────

def _detect_format(name: str, content_type: str, header: bytes) -> str:
    """
    Détecte le format du fichier avec priorité :
    magic bytes > extension > content_type.
    Retourne : 'pdf' | 'docx' | 'doc' | 'odt' | 'rtf' | 'txt' |
               'jpg' | 'png' | 'webp' | 'tiff' | 'bmp' | 'unknown'
    """
    # 1. Magic bytes (le plus fiable)
    for magic, fmt in MAGIC_BYTES:
        if header.startswith(magic):
            # Différencier docx et odt (tous deux ZIP)
            if fmt == "docx":
                name_lower = name.lower()
                if name_lower.endswith(".odt"):
                    return "odt"
            return fmt

    # 2. Extension
    ext = (name or "").lower().rsplit(".", 1)[-1]
    ext_map = {
        "pdf": "pdf", "docx": "docx", "doc": "doc",
        "odt": "odt", "rtf": "rtf", "txt": "txt",
        "jpg": "jpg", "jpeg": "jpg", "png": "png",
        "webp": "webp", "tiff": "tiff", "tif": "tiff", "bmp": "bmp",
    }
    if ext in ext_map:
        return ext_map[ext]

    # 3. Content-type
    ct = (content_type or "").lower()
    if "pdf" in ct:
        return "pdf"
    if ct in DOCX_MIME or "word" in ct or "document" in ct:
        return "docx"
    if "odt" in ct or "opendocument" in ct:
        return "odt"
    if "rtf" in ct:
        return "rtf"
    if ct == "text/plain":
        return "txt"
    if ct in IMAGE_MIME or "image" in ct:
        ext_img = {"jpeg": "jpg", "jpg": "jpg", "png": "png", "webp": "webp", "tiff": "tiff", "bmp": "bmp"}
        for k, v in ext_img.items():
            if k in ct:
                return v
        return "jpg"

    return "unknown"


def _read_bytes(source: Any) -> bytes:
    """Lit les bytes depuis un fichier-like, bytes, ou UploadedFile Django."""
    if isinstance(source, (bytes, bytearray)):
        return bytes(source)
    if hasattr(source, "read"):
        if hasattr(source, "seek"):
            source.seek(0)
        return source.read()
    return b""


def _to_buffer(source: Any) -> BytesIO:
    return BytesIO(_read_bytes(source))


# ─────────────────────────────────────────────────────────────
# NETTOYAGE TEXTE
# ─────────────────────────────────────────────────────────────

def _clean_text(raw: str) -> str:
    """
    Nettoie et normalise le texte extrait :
    - Normalisation Unicode (NFC)
    - Suppression des caractères de contrôle (sauf newline/tab)
    - Collapse des espaces multiples et lignes vides excessives
    - Suppression des artefacts PDF courants (numéros de page isolés, etc.)
    """
    if not raw:
        return ""

    # Normalisation Unicode
    text = unicodedata.normalize("NFC", raw)

    # Suppression des caractères de contrôle non imprimables
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)

    # Espaces insécables et variantes → espace ordinaire
    text = text.replace("\u00a0", " ").replace("\u200b", "").replace("\ufeff", "")

    # Tirets longs → tiret simple
    text = text.replace("\u2013", "-").replace("\u2014", "-")

    # Artefacts PDF : lignes ne contenant qu'un numéro de page
    text = re.sub(r"(?m)^\s*\d{1,3}\s*$", "", text)

    # Collapse lignes vides multiples → double saut max
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse espaces multiples sur une ligne
    text = re.sub(r"[ \t]{2,}", " ", text)

    # Suppression des lignes ne contenant que des tirets/étoiles (séparateurs PDF)
    text = re.sub(r"(?m)^[\s\-=_*•·]{3,}\s*$", "", text)

    return text.strip()


def _quality_score(text: str) -> float:
    """
    Calcule un score de qualité 0.0–1.0 basé sur :
    - Ratio lettres/total caractères (texte réel vs artefacts)
    - Longueur du texte (texte court → faible score)
    - Absence de caractères illisibles (□, ?, séquences hex)
    """
    if not text:
        return 0.0
    total = len(text)
    if total < 20:
        return 0.0

    letter_ratio = sum(1 for c in text if c.isalpha()) / total
    length_bonus = min(1.0, total / 2000)  # plafonné à 1.0 à partir de 2000 chars

    # Pénalité pour caractères illisibles
    garbage_count = len(re.findall(r"[\ufffd\u25a1]|(\?{3,})", text))
    garbage_ratio = garbage_count / max(total, 1)
    garbage_penalty = max(0.0, 1.0 - garbage_ratio * 10)

    score = (letter_ratio * 0.5 + length_bonus * 0.3 + garbage_penalty * 0.2)
    return round(min(score, 1.0), 3)


# ─────────────────────────────────────────────────────────────
# EXTRACTEURS PAR FORMAT
# ─────────────────────────────────────────────────────────────

def _extract_pdf_pypdf(buf: BytesIO) -> tuple[str, int]:
    """Extraction PDF via pypdf. Retourne (texte, nb_pages)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf non installé : pip install pypdf")
    buf.seek(0)
    reader = PdfReader(buf)
    pages = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            pages.append(t)
    return "\n".join(pages), len(reader.pages)


def _extract_pdf_pdfminer(buf: BytesIO) -> str:
    """Extraction PDF via pdfminer.six (meilleur sur PDF complexes/colonnes)."""
    try:
        from pdfminer.high_level import extract_text as pm_extract
    except ImportError:
        raise ImportError("pdfminer.six non installé : pip install pdfminer.six")
    buf.seek(0)
    return pm_extract(buf) or ""


def _extract_pdf_ocr(buf: BytesIO) -> tuple[str, int]:
    """
    OCR sur PDF scanné via pdf2image + pytesseract.
    Chaque page est rastérisée puis passée à Tesseract.
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        raise ImportError(
            "OCR non disponible. Installez : pip install pdf2image pytesseract\n"
            "et Tesseract : https://tesseract-ocr.github.io/tessdoc/Installation.html"
        )
    buf.seek(0)
    images = convert_from_bytes(buf.read(), dpi=200)
    parts = []
    for img in images:
        t = pytesseract.image_to_string(img, lang="fra+eng")
        if t.strip():
            parts.append(t)
    return "\n".join(parts), len(images)


def _extract_docx(buf: BytesIO) -> tuple[str, dict]:
    """
    Extraction Word (.docx) complète :
    paragraphes + tableaux + zones de texte + en-têtes/pieds de page.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx non installé : pip install python-docx")

    buf.seek(0)
    doc = Document(buf)
    parts = []
    meta = {}

    # Propriétés du document
    try:
        cp = doc.core_properties
        meta = {
            "author": cp.author or "",
            "created": str(cp.created or ""),
            "modified": str(cp.modified or ""),
        }
    except Exception:
        pass

    # Paragraphes principaux
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    # En-têtes et pieds de page
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container:
                for para in container.paragraphs:
                    t = para.text.strip()
                    if t and len(t) > 3:
                        parts.append(t)

    # Zones de texte (text boxes — stockées dans les formes XML)
    try:
        for shape in doc.element.body.iter(qn("w:txbxContent")):
            for p in shape.iter(qn("w:p")):
                texts = [r.text for r in p.iter(qn("w:t")) if r.text]
                t = "".join(texts).strip()
                if t:
                    parts.append(t)
    except Exception:
        pass

    return "\n".join(parts), meta


def _extract_doc_legacy(buf: BytesIO) -> str:
    """Extraction .doc (format binaire OLE2) via antiword ou textract."""
    # Tentative via python-docx2txt (parfois compatible .doc)
    try:
        import docx2txt
        buf.seek(0)
        return docx2txt.process(buf) or ""
    except Exception:
        pass
    # Fallback : extraction de chaînes UTF-16 brutes (dernier recours)
    try:
        buf.seek(0)
        raw = buf.read()
        # Les .doc stockent du texte en UTF-16 LE
        text = raw.decode("utf-16-le", errors="ignore")
        # Filtre des caractères imprimables uniquement
        text = re.sub(r"[^\x20-\x7e\u00c0-\u024f\n]", " ", text)
        text = re.sub(r" {3,}", " ", text)
        return text[:50000]  # limite de sécurité
    except Exception:
        return ""


def _extract_odt(buf: BytesIO) -> str:
    """Extraction ODT (OpenDocument Text) via odfpy."""
    try:
        from odf.opendocument import load as odf_load
        from odf.text import P
        from odf.teletype import extractText
    except ImportError:
        raise ImportError("odfpy non installé : pip install odfpy")
    buf.seek(0)
    doc = odf_load(buf)
    parts = [extractText(p) for p in doc.getElementsByType(P) if extractText(p).strip()]
    return "\n".join(parts)


def _extract_rtf(buf: BytesIO) -> str:
    """Extraction RTF via striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        raise ImportError("striprtf non installé : pip install striprtf")
    buf.seek(0)
    raw = buf.read().decode("latin-1", errors="ignore")
    return rtf_to_text(raw) or ""


def _extract_image_ocr(buf: BytesIO) -> str:
    """OCR sur image CV (JPEG, PNG, etc.) via pytesseract."""
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        raise ImportError(
            "OCR image non disponible. Installez : pip install Pillow pytesseract\n"
            "et Tesseract : https://tesseract-ocr.github.io/tessdoc/Installation.html"
        )
    buf.seek(0)
    img = Image.open(buf)
    # Amélioration légère : conversion en niveaux de gris
    img = img.convert("L")
    return pytesseract.image_to_string(img, lang="fra+eng") or ""


# ─────────────────────────────────────────────────────────────
# PIPELINE PRINCIPAL
# ─────────────────────────────────────────────────────────────

def extract_cv(
    source: Any,
    filename: str = "",
    content_type: str = "",
    *,
    allow_ocr: bool | None = None,
) -> ExtractionResult:
    """
    Extrait le texte d'un CV avec pipeline multi-moteur + fallback automatique.

    Args:
        source      : UploadedFile Django, fichier-like (read/seek), ou bytes bruts.
        filename    : Nom du fichier (pour détection de format par extension).
        content_type: MIME type déclaré (optionnel, peut être incorrect).

    Returns:
        ExtractionResult avec .text, .method, .quality_score, .page_count, .metadata.

    Exemple :
        result = extract_cv(request.FILES["cv"], request.FILES["cv"].name)
        if result.is_sufficient:
            candidate.raw_cv_text = result.text
            candidate.cv_extraction_meta = result.to_dict()
    """
    result = ExtractionResult()

    if allow_ocr is None:
        try:
            from django.conf import settings
            allow_ocr = bool(getattr(settings, 'CV_EXTRACTION_ALLOW_OCR', False))
        except Exception:
            allow_ocr = False

    # Lecture des bytes pour la détection de format
    raw_bytes = _read_bytes(source)
    if not raw_bytes:
        result.warnings.append("Fichier vide ou illisible.")
        return result

    # Reset du pointeur après lecture
    if hasattr(source, "seek"):
        try:
            source.seek(0)
        except Exception:
            pass

    fmt = _detect_format(filename, content_type, raw_bytes[:16])
    buf = BytesIO(raw_bytes)

    logger.debug("cv_extraction v2: fichier=%s format=%s taille=%d", filename, fmt, len(raw_bytes))

    # ── PDF ───────────────────────────────────────────────────
    if fmt == "pdf":
        # Tentative 1 : pypdf
        try:
            text, pages = _extract_pdf_pypdf(buf)
            result.page_count = pages
            if len(text.strip()) >= PDF_OCR_FALLBACK_THRESHOLD:
                result.text = _clean_text(text)
                result.method = ExtractionMethod.PYPDF
                logger.debug("cv_extraction v2: pypdf ok pages=%d chars=%d", pages, len(result.text))
            else:
                result.warnings.append(f"pypdf : texte insuffisant ({len(text.strip())} chars), tentative pdfminer.")
                raise ValueError("Texte insuffisant pour pypdf")
        except Exception as e_pypdf:
            # Tentative 2 : pdfminer
            try:
                buf.seek(0)
                text = _extract_pdf_pdfminer(buf)
                if len(text.strip()) >= PDF_OCR_FALLBACK_THRESHOLD:
                    result.text = _clean_text(text)
                    result.method = ExtractionMethod.PDFMINER
                    logger.debug("cv_extraction v2: pdfminer ok chars=%d", len(result.text))
                else:
                    result.warnings.append(f"pdfminer : texte insuffisant ({len(text.strip())} chars), tentative OCR.")
                    raise ValueError("Texte insuffisant pour pdfminer")
            except Exception as e_pm:
                # Tentative 3 : OCR (PDF scanné) — optionnel (désactivé par défaut en prod Railway)
                if allow_ocr:
                    try:
                        buf.seek(0)
                        text, pages = _extract_pdf_ocr(buf)
                        result.text = _clean_text(text)
                        result.method = ExtractionMethod.OCR
                        result.ocr_used = True
                        if result.page_count == 0:
                            result.page_count = pages
                        logger.info("cv_extraction v2: OCR utilisé pour PDF scanné fichier=%s", filename)
                    except ImportError as ie:
                        result.warnings.append(f"OCR non disponible : {ie}")
                        logger.warning("cv_extraction v2: OCR indisponible pour %s", filename)
                    except Exception as e_ocr:
                        result.warnings.append(f"Toutes les méthodes PDF ont échoué : {e_ocr}")
                        logger.error("cv_extraction v2: échec total PDF fichier=%s", filename, exc_info=True)
                else:
                    result.warnings.append(
                        "Texte PDF insuffisant (pypdf/pdfminer). "
                        "Utilisez un PDF texte ou un document Word, ou activez CV_EXTRACTION_ALLOW_OCR."
                    )

    # ── DOCX ─────────────────────────────────────────────────
    elif fmt == "docx":
        try:
            text, meta = _extract_docx(buf)
            result.text = _clean_text(text)
            result.method = ExtractionMethod.DOCX
            result.metadata = meta
        except ImportError as ie:
            result.warnings.append(str(ie))
        except Exception as e:
            result.warnings.append(f"Erreur extraction docx : {e}")
            logger.warning("cv_extraction v2: erreur docx fichier=%s : %s", filename, e)

    # ── DOC (binaire legacy) ──────────────────────────────────
    elif fmt == "doc":
        try:
            text = _extract_doc_legacy(buf)
            result.text = _clean_text(text)
            result.method = ExtractionMethod.DOC
            if not result.is_sufficient:
                result.warnings.append("Format .doc legacy : extraction partielle possible.")
        except Exception as e:
            result.warnings.append(f"Erreur extraction .doc : {e}")
            logger.warning("cv_extraction v2: erreur doc fichier=%s : %s", filename, e)

    # ── ODT ──────────────────────────────────────────────────
    elif fmt == "odt":
        try:
            text = _extract_odt(buf)
            result.text = _clean_text(text)
            result.method = ExtractionMethod.ODT
        except ImportError as ie:
            result.warnings.append(str(ie))
        except Exception as e:
            result.warnings.append(f"Erreur extraction ODT : {e}")

    # ── RTF ──────────────────────────────────────────────────
    elif fmt == "rtf":
        try:
            text = _extract_rtf(buf)
            result.text = _clean_text(text)
            result.method = ExtractionMethod.RTF
        except ImportError as ie:
            result.warnings.append(str(ie))
        except Exception as e:
            result.warnings.append(f"Erreur extraction RTF : {e}")

    # ── TXT ──────────────────────────────────────────────────
    elif fmt == "txt":
        try:
            buf.seek(0)
            raw = buf.read()
            # Détection d'encodage : UTF-8 puis Latin-1
            for enc in ("utf-8-sig", "utf-8", "latin-1"):
                try:
                    text = raw.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = raw.decode("latin-1", errors="replace")
            result.text = _clean_text(text)
            result.method = ExtractionMethod.PLAIN_TEXT
        except Exception as e:
            result.warnings.append(f"Erreur lecture TXT : {e}")

    # ── Images (OCR) ─────────────────────────────────────────
    elif fmt in ("jpg", "png", "webp", "tiff", "bmp"):
        try:
            text = _extract_image_ocr(buf)
            result.text = _clean_text(text)
            result.method = ExtractionMethod.IMAGE_OCR
            result.ocr_used = True
        except ImportError as ie:
            result.warnings.append(str(ie))
        except Exception as e:
            result.warnings.append(f"Erreur OCR image : {e}")
            logger.warning("cv_extraction v2: erreur OCR image fichier=%s : %s", filename, e)

    else:
        result.warnings.append(
            f"Format non supporté : '{fmt}' (fichier={filename}, content_type={content_type}). "
            f"Formats acceptés : PDF, DOCX, DOC, ODT, RTF, TXT, JPG, PNG, WEBP, TIFF, BMP."
        )
        logger.warning("cv_extraction v2: format inconnu fichier=%s type=%s", filename, content_type)

    # ── Score qualité final ───────────────────────────────────
    result.quality_score = _quality_score(result.text)

    if result.text and not result.is_sufficient:
        result.warnings.append(
            f"Texte extrait très court ({len(result.text)} chars < {MINIMUM_TEXT_LENGTH}). "
            "Le scoring ML sera moins précis."
        )

    logger.info(
        "cv_extraction v2: fichier=%s méthode=%s chars=%d qualité=%.2f ocr=%s warnings=%d",
        filename, result.method.value, len(result.text),
        result.quality_score, result.ocr_used, len(result.warnings),
    )

    return result


# ─────────────────────────────────────────────────────────────
# COMPATIBILITÉ ASCENDANTE (remplace l'ancienne API)
# ─────────────────────────────────────────────────────────────

def extract_text_from_uploaded_file(uploaded_file: Any) -> str:
    """
    Compatibilité ascendante avec la v1.
    Préférer extract_cv() pour accéder aux métadonnées et au score qualité.
    """
    if uploaded_file is None:
        return ""
    name = getattr(uploaded_file, "name", "") or ""
    ct = getattr(uploaded_file, "content_type", "") or ""
    result = extract_cv(uploaded_file, filename=name, content_type=ct)
    # Reset pointeur pour sauvegarde ultérieure (comportement v1 préservé)
    if hasattr(uploaded_file, "seek"):
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
    return result.text